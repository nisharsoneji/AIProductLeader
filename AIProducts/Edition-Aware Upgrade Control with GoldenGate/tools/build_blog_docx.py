from pathlib import Path
import re

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT_DIR = Path(__file__).resolve().parents[1]
SOURCE_MD = PROJECT_DIR / "docs" / "edition-aware-upgrade-control-with-goldengate.md"
OUTPUT_DOCX = PROJECT_DIR / "docs" / "edition-aware-upgrade-control-with-goldengate.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), "9360")
    tbl_w.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            cell.width = widths[idx]
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "0563C1")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    run.append(text_node)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_inline_runs(paragraph, text):
    pattern = re.compile(r"(`[^`]+`)|(\[[^\]]+\]\([^)]+\))|(\*\*[^*]+\*\*)")
    position = 0
    for match in pattern.finditer(text):
        if match.start() > position:
            paragraph.add_run(text[position : match.start()])
        token = match.group(0)
        if token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(10)
        elif token.startswith("["):
            label, url = re.match(r"\[([^\]]+)\]\(([^)]+)\)", token).groups()
            add_hyperlink(paragraph, label, url)
        elif token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        position = match.end()
    if position < len(text):
        paragraph.add_run(text[position:])


def apply_styles(document):
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for style_name, size, color in [
        ("Title", 24, "0B2545"),
        ("Subtitle", 11, "555555"),
        ("Heading 1", 16, "2E74B5"),
        ("Heading 2", 13, "2E74B5"),
        ("Heading 3", 12, "1F4D78"),
    ]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        if style_name.startswith("Heading"):
            style.font.bold = True
            style.paragraph_format.space_before = Pt(12 if style_name != "Heading 1" else 16)
            style.paragraph_format.space_after = Pt(6)

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(11)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def widths_for_table(column_count):
    if column_count == 2:
        return [2800, 6560]
    if column_count == 3:
        return [2200, 2500, 4660]
    if column_count == 4:
        return [1800, 1800, 2800, 2960]
    if column_count == 5:
        return [1200, 1300, 2400, 2200, 2260]
    if column_count == 6:
        return [1050, 950, 1350, 2300, 1800, 1910]
    return [int(9360 / column_count)] * column_count


def add_markdown_table(document, rows):
    header = [cell.strip() for cell in rows[0].strip("|").split("|")]
    body_rows = rows[2:]
    table = document.add_table(rows=1, cols=len(header))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for idx, text in enumerate(header):
        add_inline_runs(hdr[idx].paragraphs[0], text)
        for run in hdr[idx].paragraphs[0].runs:
            run.bold = True
        set_cell_shading(hdr[idx], "F2F4F7")
    for row_text in body_rows:
        values = [cell.strip() for cell in row_text.strip("|").split("|")]
        row = table.add_row().cells
        for idx, text in enumerate(values):
            paragraph = row[idx].paragraphs[0]
            if idx == 0 or text.lower() in {"running", "stopped", "source", "target"}:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline_runs(paragraph, text)
    set_table_width(table, widths_for_table(len(header)))
    document.add_paragraph()


def add_code_block(document, code, language):
    paragraph = document.add_paragraph()
    if language:
        label = paragraph.add_run(language.upper())
        label.bold = True
        label.font.color.rgb = RGBColor.from_string("1F4D78")
        paragraph.add_run("\n")
    run = paragraph.add_run(code.rstrip())
    run.font.name = "Consolas"
    run.font.size = Pt(9)
    paragraph.paragraph_format.left_indent = Inches(0.2)
    paragraph.paragraph_format.right_indent = Inches(0.2)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    p_pr = paragraph._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F9")
    p_pr.append(shd)


def build_docx():
    document = Document()
    apply_styles(document)

    lines = SOURCE_MD.read_text(encoding="utf-8").splitlines()
    idx = 0
    in_code = False
    code_language = ""
    code_lines = []
    pending_table = []

    def flush_table():
        nonlocal pending_table
        if pending_table:
            add_markdown_table(document, pending_table)
            pending_table = []

    while idx < len(lines):
        line = lines[idx]
        stripped = line.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(document, "\n".join(code_lines), code_language)
                in_code = False
                code_language = ""
                code_lines = []
            else:
                flush_table()
                in_code = True
                code_language = stripped[3:].strip()
            idx += 1
            continue

        if in_code:
            code_lines.append(line)
            idx += 1
            continue

        if stripped.startswith("|") and stripped.endswith("|"):
            pending_table.append(line)
            idx += 1
            continue
        flush_table()

        if not stripped:
            idx += 1
            continue

        if stripped.startswith("# "):
            paragraph = document.add_paragraph(style="Title")
            add_inline_runs(paragraph, stripped[2:])
        elif stripped.startswith("## "):
            document.add_heading(stripped[3:], level=1)
        elif stripped.startswith("### "):
            document.add_heading(stripped[4:], level=2)
        elif stripped.startswith("- "):
            paragraph = document.add_paragraph(style="List Bullet")
            add_inline_runs(paragraph, stripped[2:])
        elif re.match(r"^\d+\. ", stripped):
            paragraph = document.add_paragraph(style="List Number")
            add_inline_runs(paragraph, re.sub(r"^\d+\. ", "", stripped))
        else:
            paragraph = document.add_paragraph()
            add_inline_runs(paragraph, stripped)
        idx += 1

    flush_table()

    footer = document.sections[0].footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer.add_run("Edition-Aware Upgrade Control with Oracle GoldenGate").font.size = Pt(9)

    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT_DOCX)


if __name__ == "__main__":
    build_docx()
